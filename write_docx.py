import os
import sys
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from collections import defaultdict

class JSONToDocxConverter:
    def __init__(self, json_path, output_dir="output_docx", font_name="Times New Roman", font_size=14):
        self.json_path = json_path
        self.output_dir = output_dir
        self.font_name = font_name
        self.font_size = font_size

        os.makedirs(self.output_dir, exist_ok=True)

    def load_json(self):
        if not os.path.exists(self.json_path):
            raise FileNotFoundError(f"JSON file {self.json_path} not found.")

        with open(self.json_path, "r", encoding="utf-8") as f:
            return json.load(f)

    def filter_by_text(self, data):
        filtered_data = [x for x in data if len(x['text']) > 1]
        for idx, element in enumerate(sorted(filtered_data, key=lambda x: x['order']), start=1):
            element['order'] = idx
        return filtered_data

    def find_stats(self, numbers):
        numbers = [n for n in numbers if isinstance(n, (int, float)) and n != ""]
        numbers.sort()
        groups = defaultdict(int)
        current_group = []

        for number in numbers:
            if not current_group or abs(number - current_group[0]) <= 20:
                current_group.append(number)
            else:
                group_range = f"{min(current_group)}-{max(current_group)}" if min(current_group) != max(current_group) else f"{min(current_group)}"
                groups[group_range] += len(current_group)
                current_group = [number]

        if current_group:
            group_range = f"{min(current_group)}-{max(current_group)}" if min(current_group) != max(current_group) else f"{min(current_group)}"
            groups[group_range] += len(current_group)

        sorted_groups = sorted(groups.items(), key=lambda x: -x[1])

        transformed_data = []
        for key, count in sorted_groups:
            try:
                if '-' in key:
                    start, end = map(int, key.split('-'))
                    transformed_data.append((list(range(start, end + 1)), count))
                else:
                    transformed_data.append(([int(key)], count))
            except ValueError:
                print(f"Skipping invalid key: {key}")
                continue

        return transformed_data


    # def find_stats(self, numbers):
    #     numbers.sort()
    #     groups = defaultdict(int)
    #     current_group = []

    #     for number in numbers:
    #         if not current_group or abs(number - current_group[0]) <= 20:
    #             current_group.append(number)
    #         else:
    #             group_range = f"{min(current_group)}-{max(current_group)}" if min(current_group) != max(current_group) else f"{min(current_group)}"
    #             groups[group_range] += len(current_group)
    #             current_group = [number]

    #     if current_group:
    #         group_range = f"{min(current_group)}-{max(current_group)}" if min(current_group) != max(current_group) else f"{min(current_group)}"
    #         groups[group_range] += len(current_group)

    #     sorted_groups = sorted(groups.items(), key=lambda x: -x[1])

    #     transformed_data = []
    #     for key, count in sorted_groups:
    #         if '-' in key:
    #             start, end = map(int, key.split('-'))
    #             transformed_data.append((list(range(start, end + 1)), count))
    #         else:
    #             transformed_data.append(([int(key)], count))

    #     return transformed_data

    def process_json(self, data):
        # First step: Write numbers to lists
        left_length = [entry['horizontal_length_left'] for entry in data]
        right_length = [entry['horizontal_length_right'] for entry in data]
        vertical_length = [entry['vertical_length_to_previous'] for entry in data]

        # for entry in data:
        #     entry['ratio'] = entry['area'] / len(entry['text'])

        # ratio_s = [int(entry['ratio']) for entry in data]
        # ratio_stats = self.find_stats(ratio_s)
        left_stats = self.find_stats(left_length)
        # right_stats = self.find_stats(right_length)
        vertical_stats = self.find_stats(vertical_length)

        # Second step: Using statistics write new elements in JSON
        data[0]['paragraph'] = 'cs'
        for entry in data[1:]:
            if entry['horizontal_length_left'] in left_stats[0][0] and entry['vertical_length_to_previous'] <= max(vertical_stats[0][0]):
                entry['paragraph'] = 'lc'
            elif entry['horizontal_length_left'] in left_stats[0][0] and entry['vertical_length_to_previous'] > max(vertical_stats[0][0]):
                entry['paragraph'] = 'ls'
            elif entry['horizontal_length_left'] in left_stats[1][0] and entry['vertical_length_to_previous'] <= max(vertical_stats[0][0]):
                entry['paragraph'] = 'ls'
            elif entry['horizontal_length_left'] not in left_stats[0][0] + left_stats[1][0] and abs(entry['horizontal_length_left'] - entry['horizontal_length_right']) < 10 and entry['vertical_length_to_previous'] <= max(vertical_stats[0][0]): #and entry['ratio'] <= max(ratio_stats[0][0]):
                entry['paragraph'] = 'cc'
            elif entry['horizontal_length_left'] not in left_stats[0][0] + left_stats[1][0] and abs(entry['horizontal_length_left'] - entry['horizontal_length_right']) < 10 and entry['vertical_length_to_previous'] <= max(vertical_stats[0][0]): # and entry['ratio'] > max(ratio_stats[0][0]):
                entry['paragraph'] = 'cs'
            elif entry['horizontal_length_left'] not in left_stats[0][0] + left_stats[1][0] and abs(entry['horizontal_length_left'] - entry['horizontal_length_right']) < 10 and entry['vertical_length_to_previous'] > max(vertical_stats[0][0]):
                entry['paragraph'] = 'cs'
            elif entry['horizontal_length_left'] not in left_stats[0][0] + left_stats[1][0] and entry['horizontal_length_left'] - entry['horizontal_length_right'] > 10 and entry['vertical_length_to_previous'] <= max(vertical_stats[0][0]):
                entry['paragraph'] = 'rc'
            elif entry['horizontal_length_left'] not in left_stats[0][0] + left_stats[1][0] and entry['horizontal_length_left'] - entry['horizontal_length_right'] > 10 and entry['vertical_length_to_previous'] > max(vertical_stats[0][0]):
                entry['paragraph'] = 'rs'
            else:
                entry['paragraph'] = 'ls'

        for entry in data:
            if max(entry["color_code"]) < 200:
                entry['font_type'] = 'bold'
            else:
                entry['font_type'] = 'normal'

        keys_to_keep = {'text', 'order', 'paragraph', 'font_type'}
        filtered_json = [{key: d[key] for key in keys_to_keep if key in d} for d in data]
        return self.process_text(filtered_json)

    def process_text(self, data):
        result = []
        previous = None

        for current in data:
            if previous:
                # Check the concatenation rules based on the 'paragraph' values
                if current['paragraph'] == 'lc' and previous['paragraph'] == 'ls':
                    previous['text'] += ' ' + current['text']
                elif current['paragraph'] == 'lc' and previous['paragraph'] == 'lc':
                    previous['text'] += ' ' + current['text']
                elif current['paragraph'] == 'cc' and previous['paragraph'] == 'cs':
                    previous['text'] += ' ' + current['text']
                elif current['paragraph'] == 'cc' and previous['paragraph'] == 'cc':
                    previous['text'] += ' ' + current['text']
                elif current['paragraph'] == 'rc' and previous['paragraph'] == 'rs':
                    previous['text'] += ' ' + current['text']
                elif current['paragraph'] == 'rc' and previous['paragraph'] == 'rc':
                    previous['text'] += ' ' + current['text']
                else:
                    # If no matching condition, push the previous element to the result
                    result.append(previous)
                    previous = current
            else:
                # For the first item, just set it as the previous item
                previous = current

        # Add the last remaining item
        if previous:
            result.append(previous)

        # Rewriting the 'order' correctly
        for idx, item in enumerate(result):
            item['order'] = idx + 1

        return result

    def write_docx(self, data, output_filename="output.docx"):
        document = Document()

        for entry in data:
            text = entry['text'].replace("\n", "")

            paragraph = document.add_paragraph()
            if entry['paragraph'] == 'cs':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif entry['paragraph'] == 'ls':
                paragraph.paragraph_format.first_line_indent = Pt(35)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            elif entry['paragraph'] == 'rs':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            run = paragraph.add_run(text)
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)
            if entry['font_type'] == 'bold':
                run.bold = True

        output_path = os.path.join(self.output_dir, output_filename)
        document.save(output_path)
        print(f"Document saved to {output_path}")

    # def convert(self):
    #     data = self.load_json()
    #     filtered_data = self.filter_by_text(data)
    #     processed_data = self.process_json(filtered_data)
    #     self.write_docx(processed_data)

    def convert(self, output_filename):
        data = self.load_json()
        filtered_data = self.filter_by_text(data)
        processed_data = self.process_json(filtered_data)
        self.write_docx(processed_data, output_filename)


    # @staticmethod
    # def run_from_command_line():
    #     if len(sys.argv) != 2:
    #         print("Usage: python write_docx.py <json_file_path>")
    #         sys.exit(1)

    #     json_file_path = sys.argv[1]
    #     try:
    #         converter = JSONToDocxConverter(json_file_path)
    #         converter.convert()
    #     except Exception as e:
    #         print(f"Error: {e}")
    @staticmethod
    def run_from_command_line():
        if len(sys.argv) != 3:
            print("Usage: python write_docx.py <json_file_path> <output_docx_path>")
            sys.exit(1)

        json_file_path = sys.argv[1]
        output_docx_path = sys.argv[2]

        try:
            converter = JSONToDocxConverter(json_file_path)
            converter.convert(output_docx_path)
        except Exception as e:
            print(f"Error: {e}")

if __name__ == "__main__":
    JSONToDocxConverter.run_from_command_line()